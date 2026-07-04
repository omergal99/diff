"""
Office document text extraction.

Extracts plain text from DOCX, XLSX, and PDF files.
Each extractor is isolated and returns a clean string.

Dependencies:
  python-docx  → DOCX
  openpyxl     → XLSX
  PyMuPDF      → PDF
"""

from __future__ import annotations
import io
from utils.logger import get_logger

log = get_logger("office")


def extract_docx(data: bytes) -> str:
    """Extract all paragraph text from a DOCX file."""
    try:
        import docx
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    text = "\n".join(p for p in paragraphs if p.strip())
    log.debug(f"DOCX: extracted {len(text)} chars, {len(paragraphs)} paragraphs")
    return text


def extract_xlsx(data: bytes) -> str:
    """
    Extract all cell values from an XLSX file.
    Returns a tab-separated, newline-separated representation of all sheets.
    """
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("openpyxl not installed. Run: pip install openpyxl")

    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    lines: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            lines.append("\t".join(cells))
    wb.close()
    text = "\n".join(lines)
    log.debug(f"XLSX: extracted {len(text)} chars from {len(wb.sheetnames)} sheets")
    return text


def extract_pdf(data: bytes) -> str:
    """Extract all text from a PDF file, preserving page breaks."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("PyMuPDF not installed. Run: pip install PyMuPDF")

    doc = fitz.open(stream=data, filetype="pdf")
    pages: list[str] = []
    for i, page in enumerate(doc, 1):
        text = page.get_text("text")
        if text.strip():
            pages.append(f"--- Page {i} ---\n{text}")
    doc.close()
    result = "\n".join(pages)
    log.debug(f"PDF: extracted {len(result)} chars from {len(pages)} pages")
    return result


def extract(data: bytes, file_type: str) -> str:
    """
    Dispatch extraction based on file type.
    file_type: "docx" | "xlsx" | "pdf"
    """
    extractors = {
        "docx": extract_docx,
        "xlsx": extract_xlsx,
        "pdf": extract_pdf,
    }
    if file_type not in extractors:
        raise ValueError(f"Unsupported file type: {file_type}. Use: {list(extractors)}")
    return extractors[file_type](data)
